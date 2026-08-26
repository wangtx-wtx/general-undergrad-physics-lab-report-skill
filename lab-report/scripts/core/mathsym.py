# -*- coding: utf-8 -*-
"""mathtext 符号 -> docx 富文本 run 转换器
把 mathtext 风格的数学表达式转成带真上下标/斜体的文本 run，天然与正文基线对齐。
支持: 普通变量(斜体)、_下标、^上标、'撇号、反斜杠greek命令、{...}分组。
"""
import re
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

GREEK = {
    'theta': '\u03b8', 'delta': '\u03b4', 'Delta': '\u0394', 'alpha': '\u03b1',
    'beta': '\u03b2', 'gamma': '\u03b3', 'sigma': '\u03c3', 'mu': '\u03bc',
    'pi': '\u03c0', 'omega': '\u03c9', 'rho': '\u03c1', 'lambda': '\u03bb',
}

# 数学函数/算子：输出为正体
FUNCTIONS = {'ln', 'log', 'sin', 'cos', 'tan', 'exp', 'lim', 'max', 'min', 'Delta'}

def _add_run(p, text, size, italic=False, sub=False, sup=False):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.italic = italic
    if sub: r.font.subscript = True
    if sup: r.font.superscript = True
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return r

def _push(p, s, size, italic=False, sub=False, sup=False):
    if s:
        _add_run(p, s, size, italic=italic, sub=sub, sup=sup)

def math_to_runs(p, expr, size=12):
    """把 mathtext 表达式 expr 解析为带格式的 run，追加到段落 p。"""
    i = 0
    n = len(expr)
    buf = []

    def flush():
        nonlocal buf
        if buf:
            _push(p, ''.join(buf), size, italic=True)
            buf = []

    while i < n:
        c = expr[i]
        if c == '\\':
            flush()
            # 处理带参数的命令: \mathrm{...}
            if expr[i:i+7] == '\\mathrm' :
                i += 7
                # 跳过可能的空格
                while i < n and expr[i] == ' ': i += 1
                if i < n and expr[i] == '{':
                    j = expr.index('}', i)
                    inner = expr[i+1:j]
                    # 处理 \cdot, \pm 等嵌套
                    inner = inner.replace('\\cdot', '\u00b7').replace('\\pm', '\u00b1').replace('\\,', '')
                    # 去掉命令符号后残留的分隔空格，如 "g· K" -> "g·K"
                    inner = re.sub(r'([\u00b7\u00b1\u2248\u00d7])\s+', r'\1', inner)
                    _push(p, inner, size, italic=False)
                    i = j + 1
                continue
            m = re.match(r'\\([A-Za-z]+)', expr[i:])
            if m:
                name = m.group(1)
                if name == 'Delta':
                    _push(p, '\u0394', size, italic=False)
                elif name == 'ln':
                    _push(p, 'ln', size, italic=False)
                elif name == 'cdot':
                    _push(p, '\u00b7', size, italic=False)
                elif name == 'pm':
                    _push(p, '\u00b1', size, italic=False)
                elif name == 'approx':
                    _push(p, '\u2248', size, italic=False)
                elif name == 'sqrt':
                    # 根号：\sqrt{...} → √(...)，内容递归处理
                    _push(p, '\u221a', size, italic=False)
                    j = expr.find('}', i + m.end())
                    if j != -1 and expr[i + m.end()] == '{':
                        inner = expr[i + m.end() + 1:j]
                        _push(p, '(', size, italic=False)
                        math_to_runs(p, inner, size)
                        _push(p, ')', size, italic=False)
                        i = j + 1
                        continue
                    else:
                        i += m.end()
                        continue
                elif name == 'sim':
                    _push(p, '~', size, italic=False)
                elif name == 'times':
                    _push(p, '\u00d7', size, italic=False)
                elif name == 'le':
                    _push(p, '\u2264', size, italic=False)
                elif name == 'ge':
                    _push(p, '\u2265', size, italic=False)
                elif name in GREEK:
                    _push(p, GREEK[name], size, italic=True)
                elif name in FUNCTIONS:
                    _push(p, name, size, italic=False)
                else:
                    _push(p, name, size, italic=True)
                i += m.end()
            else:
                # 非字母命令，如 \   (空格)
                i += 1
                if i < n and expr[i] == ' ':
                    _push(p, ' ', size, italic=False)
                    i += 1
            continue
        if c == '_':
            flush(); i += 1
            if i < n and expr[i] == '{':
                j = expr.index('}', i)
                _push(p, expr[i+1:j], size, sub=True)
                i = j + 1
            else:
                if i < n:
                    _push(p, expr[i], size, sub=True); i += 1
            continue
        if c == '^':
            flush(); i += 1
            if i < n and expr[i] == '{':
                j = expr.index('}', i)
                _push(p, expr[i+1:j], size, sup=True)
                i = j + 1
            else:
                if i < n:
                    _push(p, expr[i], size, sup=True); i += 1
            continue
        if c == "'":
            flush(); _push(p, '\u2032', size, sup=True); i += 1
            continue
        if c in ' \t,;':
            flush(); i += 1; continue
        if c in '()-=+':
            flush(); _push(p, c, size, italic=False); i += 1; continue
        if c.isalpha():
            buf.append(c); i += 1; continue
        else:
            flush(); _push(p, c, size, italic=False); i += 1; continue
    flush()


if __name__ == '__main__':
    import os, tempfile
    from docx import Document
    doc = Document()
    tests = [
        r'\theta_0', r'\theta', r'm_1', r'm_1\'', r"S'", r"S''",
        r'c_1\'', r'\delta C\'', r'k/C', r'\ln(\theta-\theta_0)',
        r'S\'/S\'\'', r'\Delta\theta=\theta-\theta_0',
    ]
    for t in tests:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run('前文 '); r.font.size = Pt(12)
        math_to_runs(p, t, 12)
        r2 = p.add_run(' 后文'); r2.font.size = Pt(12)
    doc.save(os.path.join(tempfile.gettempdir(), '_mathsym_test.docx'))
    print('saved mathsym test')
