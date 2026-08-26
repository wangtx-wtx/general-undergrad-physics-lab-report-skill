# -*- coding: utf-8 -*-
"""
通用 docx 报告排版帮助函数。
排版规范：汉字宋体，英文/数字 Times New Roman，
正文数学符号用 $...$ 标记（由 mathsym 渲染为真上下标文本），
独立公式用 mathtext 图片（居中）。
"""
import os, re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .mathsym import math_to_runs

CN = '宋体'
EN = 'Times New Roman'
BODY_SIZE = 12   # 小四

def set_run(run, size=BODY_SIZE, bold=False):
    """设置 run：中文宋体、英文/数字 Times New Roman。"""
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), EN)
    rFonts.set(qn('w:hAnsi'), EN)
    rFonts.set(qn('w:eastAsia'), CN)

def add_heading(doc, text, size=14):
    """章节标题（加粗黑体/宋体）。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.name = CN
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), CN); rFonts.set(qn('w:eastAsia'), CN); rFonts.set(qn('w:hAnsi'), CN)
    return p

def add_para(doc, text, size=BODY_SIZE, bold=False, align=None, indent_first=True):
    """正文段落，支持 $...$ 数学片段（真上下标，基线对齐）。"""
    from .formula import add_rich_para
    return add_rich_para(doc, text, size=size, align=align, indent_first=indent_first,
                         math_h=0.235)

def add_plain_para(doc, text, size=BODY_SIZE, bold=False, align=None, indent_first=True):
    """纯文本段落（无数学片段）。"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    if indent_first:
        pf.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold)
    return p

def add_formula(doc, key, formula_dir=None, height_in=0.44, max_width_in=5.6):
    """插入独立公式图片（水平居中）。key 为公式 png 文件名（不含扩展名）。"""
    from PIL import Image
    if formula_dir is None:
        formula_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', 'charts', 'formulas')
    fp = os.path.join(formula_dir, key + '.png')
    if not os.path.exists(fp):
        add_plain_para(doc, '[' + key + ']', indent_first=False)
        return
    with Image.open(fp) as im:
        w, h = im.size
    width_in = height_in * (w / h)
    if width_in > max_width_in:
        width_in = max_width_in
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4); pf.space_after = Pt(4); pf.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(fp, width=Inches(width_in))
    return p

def add_picture_center(doc, img_path, width_in=5.4, caption=None):
    """插入图片（居中），可选图题。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        set_run(r, size=10, bold=False)
    return p

def set_cell_text(cell, text, size=10, bold=False, center=True):
    """设置表格单元格文本（自动转字符串）。"""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.0
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for rm in list(p.runs):
        rm._element.getparent().remove(rm._element)
    run = p.add_run(str(text))
    set_run(run, size=size, bold=bold)

def add_data_table(doc, header_row, data_rows, col_widths=None, style='Table Grid'):
    """通用数据表格。header_row: 表头列表; data_rows: 二维列表。"""
    n_cols = len(header_row)
    tbl = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    tbl.style = style
    for j, h in enumerate(header_row):
        set_cell_text(tbl.cell(0, j), str(h), size=10, bold=True)
    for i, row in enumerate(data_rows):
        for j, v in enumerate(row):
            set_cell_text(tbl.cell(i + 1, j), str(v), size=10)
    return tbl

def new_document():
    """创建带标准页边距的新文档。"""
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1); s.bottom_margin = Inches(1)
        s.left_margin = Inches(1); s.right_margin = Inches(1)
    return doc
